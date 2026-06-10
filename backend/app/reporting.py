from __future__ import annotations

from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .models import ReviewTask


def render_markdown_report(task: ReviewTask) -> str:
    created = task.created_at.strftime("%Y-%m-%d %H:%M:%S")
    lines = [
        f"# CodeMate Review Report - {task.project_name}",
        "",
        f"- Task ID: {task.id}",
        f"- Language: {task.language}",
        f"- Status: {task.status}",
        f"- Created: {created}",
        f"- Source: {task.source_kind}",
        "",
        "## Summary",
        "",
        task.summary or "No summary available.",
        "",
        "## Issues",
        "",
    ]
    if task.issues:
        for issue in task.issues:
            lines.extend(
                [
                    f"### {issue.severity.upper()} - {issue.type}",
                    "",
                    f"- Location: `{issue.file}:{issue.line_start}-{issue.line_end}`",
                    f"- Confidence: {issue.confidence:.2f}",
                    f"- Status: {issue.status}",
                    f"- Source: {issue.source}",
                    "",
                    issue.description,
                    "",
                    f"Suggestion: {issue.suggestion}",
                    "",
                ]
            )
    else:
        lines.extend(["No issues recorded.", ""])

    lines.extend(["## Test Cases", ""])
    if task.test_cases:
        for case in task.test_cases:
            lines.extend(
                [
                    f"### {case.name}",
                    "",
                    f"- Category: {case.category}",
                    f"- Priority: {case.priority}",
                    f"- Status: {case.status}",
                    f"- Input: {case.input}",
                    f"- Expected: {case.expected}",
                    "",
                ]
            )
    else:
        lines.extend(["No test cases recorded.", ""])

    return "\n".join(lines)


def render_docx_report(task: ReviewTask) -> bytes:
    document = Document()
    document.add_heading(f"CodeMate Review Report - {task.project_name}", 0)
    document.add_paragraph(f"Task ID: {task.id}")
    document.add_paragraph(f"Language: {task.language}")
    document.add_paragraph(f"Status: {task.status}")
    document.add_paragraph(f"Source: {task.source_kind}")
    document.add_heading("Summary", level=1)
    document.add_paragraph(task.summary or "No summary available.")

    document.add_heading("Issues", level=1)
    if task.issues:
        for issue in task.issues:
            document.add_heading(f"{issue.severity.upper()} - {issue.type}", level=2)
            document.add_paragraph(
                f"Location: {issue.file}:{issue.line_start}-{issue.line_end}"
            )
            document.add_paragraph(f"Confidence: {issue.confidence:.2f}")
            document.add_paragraph(f"Status: {issue.status}")
            document.add_paragraph(f"Source: {issue.source}")
            document.add_paragraph(issue.description)
            document.add_paragraph(f"Suggestion: {issue.suggestion}")
    else:
        document.add_paragraph("No issues recorded.")

    document.add_heading("Test Cases", level=1)
    if task.test_cases:
        for case in task.test_cases:
            document.add_heading(case.name, level=2)
            document.add_paragraph(f"Category: {case.category}")
            document.add_paragraph(f"Priority: {case.priority}")
            document.add_paragraph(f"Status: {case.status}")
            document.add_paragraph(f"Input: {case.input}")
            document.add_paragraph(f"Expected: {case.expected}")
    else:
        document.add_paragraph("No test cases recorded.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf_report(task: ReviewTask) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title=f"CodeMate Task {task.id}")
    styles = getSampleStyleSheet()
    story = [
        Paragraph(f"CodeMate Review Report - {_escape(task.project_name)}", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Task ID: {task.id}", styles["Normal"]),
        Paragraph(f"Language: {_escape(task.language)}", styles["Normal"]),
        Paragraph(f"Status: {_escape(task.status)}", styles["Normal"]),
        Paragraph(f"Source: {_escape(task.source_kind)}", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("Summary", styles["Heading1"]),
        Paragraph(_escape(task.summary or "No summary available."), styles["BodyText"]),
        Spacer(1, 12),
        Paragraph("Issues", styles["Heading1"]),
    ]

    if task.issues:
        for issue in task.issues:
            story.extend(
                [
                    Paragraph(
                        f"{issue.severity.upper()} - {_escape(issue.type)}",
                        styles["Heading2"],
                    ),
                    Paragraph(
                        f"Location: {_escape(issue.file)}:{issue.line_start}-{issue.line_end}",
                        styles["Normal"],
                    ),
                    Paragraph(f"Confidence: {issue.confidence:.2f}", styles["Normal"]),
                    Paragraph(f"Status: {_escape(issue.status)}", styles["Normal"]),
                    Paragraph(f"Source: {_escape(issue.source)}", styles["Normal"]),
                    Paragraph(_escape(issue.description), styles["BodyText"]),
                    Paragraph(f"Suggestion: {_escape(issue.suggestion)}", styles["BodyText"]),
                    Spacer(1, 8),
                ]
            )
    else:
        story.append(Paragraph("No issues recorded.", styles["Normal"]))

    story.append(Paragraph("Test Cases", styles["Heading1"]))
    if task.test_cases:
        for case in task.test_cases:
            story.extend(
                [
                    Paragraph(_escape(case.name), styles["Heading2"]),
                    Paragraph(f"Category: {_escape(case.category)}", styles["Normal"]),
                    Paragraph(f"Priority: {_escape(case.priority)}", styles["Normal"]),
                    Paragraph(f"Status: {_escape(case.status)}", styles["Normal"]),
                    Paragraph(f"Input: {_escape(case.input)}", styles["BodyText"]),
                    Paragraph(f"Expected: {_escape(case.expected)}", styles["BodyText"]),
                    Spacer(1, 8),
                ]
            )
    else:
        story.append(Paragraph("No test cases recorded.", styles["Normal"]))

    doc.build(story)
    return buffer.getvalue()


def _escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
